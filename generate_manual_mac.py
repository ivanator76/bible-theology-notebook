#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate 聖經神學筆記 macOS 本機版使用手冊 .docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ── Page margins (A4) ─────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Inches(8.27)
section.page_height   = Inches(11.69)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
ACCENT   = RGBColor(0x2C, 0x5F, 0x8A)
ACCENT2  = RGBColor(0x4A, 0x7E, 0xB5)
MUTED    = RGBColor(0x55, 0x55, 0x55)
GREY     = RGBColor(0x88, 0x88, 0x88)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = ACCENT
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = ACCENT2
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run('提示：' + text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = MUTED
    return p

def sp():
    doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════════════════════
t = doc.add_heading('聖經神學筆記', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

p = doc.add_paragraph('macOS 本機版　使用手冊')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(14)
p.runs[0].font.color.rgb = MUTED

p2 = doc.add_paragraph('版本 2.0')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.runs[0].font.size = Pt(11)
p2.runs[0].font.color.rgb = GREY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 目錄
# ══════════════════════════════════════════════════════════════════════════════
h1('目錄')
for item in [
    '1.  安裝',
    '2.  首次啟動',
    '3.  介面說明',
    '4.  建立與管理筆記',
    '5.  標籤系統',
    '6.  主題鏈結',
    '7.  經文查閱',
    '8.  AI 輔助分析',
    '9.  資源庫',
    '10. 備份與還原',
    '11. 常見問題',
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(item)
    run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. 安裝
# ══════════════════════════════════════════════════════════════════════════════
h1('1. 安裝')

h2('1.1 系統需求')
bullet('macOS 12 Monterey 以上（Intel 或 Apple Silicon）')
bullet('磁碟空間：約 350 MB（含離線聖經資料）')
bullet('網路連線：僅 AI 功能需要，聖經經文查閱完全離線可用')
sp()

h2('1.2 安裝步驟')
body('① 雙擊 .dmg 安裝檔，等待掛載完成。')
body('② 將「聖經神學筆記.app」圖示拖入右側「應用程式」資料夾捷徑。')
body('③ 退出 .dmg（從桌面拖到垃圾桶，或按 ⌘E）。')
sp()

h2('1.3 移除 Gatekeeper 限制（首次必做）')
body('由於 App 未經 Apple 公證，macOS 可能顯示「已損壞」或「無法驗證開發者」的警告。請按照以下步驟解除：')
body('① 開啟「終端機」（Terminal）應用程式。')
body('② 貼上並執行以下指令：')
code('xattr -cr /Applications/聖經神學筆記.app')
body('③ 回到「應用程式」資料夾，重新雙擊開啟 App。')
tip('此步驟只需在首次安裝後執行一次。更新 App 後若再度出現警告，請重複執行。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 2. 首次啟動
# ══════════════════════════════════════════════════════════════════════════════
h1('2. 首次啟動')
body('App 啟動後會在背景自動建立資料庫，通常在 3 秒內完成並顯示主畫面。')
body('所有資料儲存於本機，路徑為：')
code('~/Library/Application Support/聖經神學筆記/data/')
body('此資料夾由系統自動建立，無需手動操作。')
tip('若要在多台電腦間同步資料，可使用「備份 / 還原」功能手動傳輸，或將上述資料夾放入 iCloud Drive 並建立符號連結（進階用法）。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 3. 介面說明
# ══════════════════════════════════════════════════════════════════════════════
h1('3. 介面說明')
body('主介面分為左右兩個區域：')
sp()

h2('左側：導覽側欄')
bullet('全部筆記：列出所有筆記，可依標籤篩選')
bullet('依書卷瀏覽：以聖經書卷分類瀏覽')
bullet('主題鏈結：跨段落的主題研究鏈管理')
bullet('資源庫：外部學習資源管理')
bullet('標籤管理：建立與管理聖經神學及系統神學標籤')
bullet('備份 / 還原：資料匯出與匯入')
sp()

h2('右側：主內容區域')
body('依據左側所選項目，顯示筆記列表、筆記詳情、編輯表單或管理頁面。')
sp()

h2('頂部搜尋列')
body('可依書卷名稱、標題或筆記內容關鍵字即時搜尋，支援中文與英文。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 4. 建立與管理筆記
# ══════════════════════════════════════════════════════════════════════════════
h1('4. 建立與管理筆記')

h2('4.1 新增筆記')
body('點擊右上角「＋ 新增筆記」，進入編輯頁面，填寫以下欄位：')
sp()

h3('書卷與章節')
bullet('書卷：從下拉選單選取（舊約 / 新約）')
bullet('起始章 / 結束章：輸入章次（單章筆記只填起始章即可）')
bullet('起始節 / 結束節（選填）：精確標記段落節次，例如 1:1–18')
sp()

h3('標題與內容')
bullet('標題（選填）：筆記的簡短名稱')
bullet('內容：支援 Markdown 語法，輸入時右側即時預覽排版')
sp()

h3('常用 Markdown 語法')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
lines = [
    '# 大標題    ## 中標題    ### 小標題',
    '**粗體**    *斜體*    > 引用區塊',
    '- 項目列表    1. 編號列表',
    '`行內程式碼`    ```多行程式碼區塊```',
]
for l in lines:
    rr = doc.add_paragraph()
    rr.paragraph_format.left_indent = Cm(1)
    run = rr.add_run(l)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
sp()

h3('標籤')
bullet('聖經神學標籤：點擊下拉選取，或直接輸入新標籤名稱後按 Enter 建立')
bullet('系統神學標籤：操作方式相同')
sp()

h2('4.2 閱讀筆記')
body('在筆記列表點擊任一筆記進入閱讀模式，顯示：')
bullet('完整筆記內容（Markdown 渲染後）')
bullet('所有標籤')
bullet('所屬主題鏈結')
bullet('連結的外部資源')
bullet('下方的經文查閱面板')
bullet('底部的 AI 輔助分析面板')
sp()

h2('4.3 編輯與刪除')
bullet('閱讀頁面右上角點擊「編輯」進入編輯模式')
bullet('點擊「刪除」並在確認對話框按下確認，即永久刪除')
tip('刪除後無法復原，建議刪除前先做備份。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 5. 標籤系統
# ══════════════════════════════════════════════════════════════════════════════
h1('5. 標籤系統')
body('本系統使用兩種獨立的標籤，讓您從不同角度標記同一段聖經文本。')
sp()

h2('5.1 聖經神學標籤（BT Tags）')
body('用於標記段落在聖經正典敘事中的主題位置與神學功能，例如：')
bullet('盟約、應許、預表、成全、王權、聖殿、出埃及、新創造……')
body('管理路徑：側欄「標籤管理」→「聖經神學標籤」')
sp()

h2('5.2 系統神學標籤（ST Tags）')
body('用於標記段落所涉及的系統神學教義分類，例如：')
bullet('神論、基督論、聖靈論、人論、罪論、救贖論、教會論、末世論……')
body('管理路徑：側欄「標籤管理」→「系統神學標籤」')
sp()

h2('5.3 透過標籤篩選筆記')
body('在「全部筆記」頁面，點擊左側任一標籤，即可篩選出包含該標籤的所有筆記。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 6. 主題鏈結
# ══════════════════════════════════════════════════════════════════════════════
h1('6. 主題鏈結（Theme Chains）')
body('主題鏈結讓您將跨書卷、跨章節的相關段落串成一條研究鏈，追蹤某個神學主題在正典中的發展軌跡。')
sp()

h2('6.1 建立主題鏈')
bullet('點擊側欄「主題鏈結」')
bullet('點擊「＋ 新增」，輸入鏈結名稱與說明')
sp()

h2('6.2 加入節點')
bullet('進入鏈結詳細頁面，點擊「加入筆記」')
bullet('搜尋並選取已有的筆記作為節點')
bullet('可拖動節點排序，以反映主題的發展順序')
sp()

h2('6.3 從筆記加入鏈結')
body('在筆記閱讀頁面，可直接將該筆記加入某條已有的主題鏈，無需切換頁面。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 7. 經文查閱
# ══════════════════════════════════════════════════════════════════════════════
h1('7. 經文查閱')
body('在筆記閱讀及編輯頁面，下方會顯示對應章節的「經文參考」面板。所有經文資料皆已離線內建於 App 中，無需網路連線即可使用。')
sp()

h2('7.1 語言切換')
body('面板頂部有四個語言按鈕：')

# Table for language buttons
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '按鈕'
hdr[1].text = '說明'
rows_data = [
    ('中文', '顯示中文和合本（CUV）'),
    ('中英', '同時顯示中文與英文（WEB）雙語對照'),
    ('English', '僅顯示英文 World English Bible（WEB）'),
    ('原文', '新約顯示希臘文（Tischendorf 8th ed.）；舊約顯示希伯來文（Westminster Leningrad Codex）'),
]
for i, (btn, desc) in enumerate(rows_data):
    row = table.rows[i+1]
    row.cells[0].text = btn
    row.cells[1].text = desc
sp()

h2('7.2 原文顯示')
bullet('新約（馬太福音至啟示錄）：顯示帶有氣息符號與重音的古希臘文（Tischendorf 第八版）')
bullet('舊約（創世記至瑪拉基書）：顯示帶有母音符號（尼庫達）的希伯來文（Westminster Leningrad Codex）')
bullet('希伯來文採由右至左（RTL）排列，為正確閱讀方向')
tip('原文版本為學術版本，保留完整的母音符號，適合配合原文工具書使用。')
sp()

h2('7.3 顯示範圍')
bullet('預設顯示筆記所指定的節次範圍（例如 1:1–18），指定節次會以金色色塊標示')
bullet('點擊面板底部「顯示整章」，可展開查看完整章節')
bullet('再次點擊「只顯示選擇節數」，可收回至指定範圍')
sp()

h2('7.4 調整字體大小')
body('面板右上角有字體大小控制器（A− / 數字 / A+），可自由調整顯示大小（範圍：11–22px）。調整後的設定會自動記憶，下次開啟 App 時保持不變。')
sp()

h2('7.5 摺疊面板')
body('點擊面板標題列的箭頭圖示可摺疊 / 展開面板，不影響筆記內容。')
sp()

tip('所有聖經資料（中文、英文、希臘文、希伯來文）均已預先內建於 App 中，不需網路連線即可查閱。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 8. AI 輔助分析
# ══════════════════════════════════════════════════════════════════════════════
h1('8. AI 輔助分析')
body('AI 功能支援三大 AI 服務供應商，提供三種智慧分析功能。使用前需設定 API Key 及網路連線。')
sp()

h2('8.1 支援的 AI 供應商')

table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = '供應商'
hdr2[1].text = '模型'
hdr2[2].text = 'API Key 格式'
providers = [
    ('Anthropic', 'Claude Sonnet（最新版）', 'sk-ant-...'),
    ('OpenAI', 'GPT-4o', 'sk-...'),
    ('Google', 'Gemini 2.0 Flash', 'AIza...'),
]
for i, (name, model, fmt) in enumerate(providers):
    row = table2.rows[i+1]
    row.cells[0].text = name
    row.cells[1].text = model
    row.cells[2].text = fmt
sp()

h2('8.2 設定 API Key')
body('① 在任一筆記閱讀頁面，找到底部「AI 輔助分析」卡片。')
body('② 點擊卡片右上角「設定 API Key」按鈕，開啟設定視窗。')
body('③ 視窗頂部有三個分頁（Anthropic / OpenAI / Google），選擇您要使用的供應商。')
body('④ 取得對應的 API Key（詳見下方連結說明），貼入輸入欄，點擊「儲存」。')
body('⑤ 儲存後畫面會顯示遮罩版本（如：sk-ant-api03-...xxxx），表示設定成功。')
body('⑥ 若已設定多個供應商，點擊「切換至此供應商」可在不同 AI 服務間切換。')
sp()

h3('取得 API Key 的網址')
bullet('Anthropic：前往 console.anthropic.com 註冊並建立 Key')
bullet('OpenAI：前往 platform.openai.com/api-keys 建立 Key')
bullet('Google：前往 aistudio.google.com 建立 Key')
sp()

tip('API Key 儲存於本機（~/Library/Application Support/聖經神學筆記/data/settings.json），不會上傳至任何外部伺服器。')
sp()

h2('8.3 三種分析功能')

h3('① 建議相關經文')
body('AI 根據筆記內容、標題及標籤，建議 5–8 段在主題、神學概念或正典脈絡上與此段落相關的經文，並附上關聯原因說明。')
sp()

h3('② 建議教義連結')
body('AI 分析此段落如何貢獻於系統神學各教義。優先從您現有的系統神學標籤中選取最相關的 2–4 個教義，說明具體的神學連結點及此段落的獨特貢獻。')
sp()

h3('③ 延伸研究方向')
body('AI 建議 3–5 個值得深入探索的研究方向，涵蓋正典關聯、歷史文化背景、跨文本比較及重要神學問題，並附上具體的研究切入點。')
sp()

h2('8.4 使用流程')
body('① 開啟任一筆記的閱讀頁面，滾動至底部的「AI 輔助分析」卡片。')
body('② 卡片標題旁會顯示目前使用的 AI 供應商名稱（如「Anthropic」）。')
body('③ 點擊三個功能按鈕之一（「建議相關經文」、「建議教義連結」、「延伸研究方向」）。')
body('④ 等待 AI 回應（通常 10–30 秒），結果將以格式化文字顯示於卡片下方。')
body('⑤ 同一次 Session 內可重複點擊不同按鈕，結果會分別保留。')
sp()
tip('AI 分析結果不會自動儲存至筆記，如需保留，請手動複製至筆記內容欄。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 9. 資源庫
# ══════════════════════════════════════════════════════════════════════════════
h1('9. 資源庫')
body('資源庫讓您集中管理與聖經研究相關的外部資源，並可連結至特定筆記。')
sp()

h2('9.1 新增資源')
bullet('點擊側欄「資源庫」，再點擊「＋ 新增資源」')
bullet('填寫資源名稱、類型（書籍 / 文章 / 影片 / 講道 / 網頁等）')
bullet('可輸入作者、出版年份、網址連結等資訊')
sp()

h2('9.2 連結資源至筆記')
bullet('在筆記閱讀或編輯頁面，找到「相關資源」區塊')
bullet('點擊「連結資源」，從資源庫中選取並加入')
bullet('可加上頁碼或段落備註，方便日後查閱')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 10. 備份與還原
# ══════════════════════════════════════════════════════════════════════════════
h1('10. 備份與還原')
body('強烈建議定期備份，以防止資料遺失。')
sp()

h2('10.1 匯出備份')
bullet('點擊側欄「備份 / 還原」')
bullet('點擊「匯出備份」，系統將下載一份 JSON 格式的完整備份檔')
bullet('備份檔包含所有筆記、標籤、主題鏈結及資源資料')
tip('建議每週或每月定期備份，並將備份檔存至 iCloud、外接硬碟或其他安全位置。')
sp()

h2('10.2 還原備份')
bullet('點擊「還原備份」，選取先前匯出的 .json 備份檔')
bullet('系統將以備份資料完全取代目前所有資料')
tip('還原操作不可逆，執行前請確認已了解目前資料將被覆蓋。')
sp()

h2('10.3 搬遷至新電腦')
body('① 在舊電腦匯出備份檔（.json）。')
body('② 在新電腦安裝 App 並完成首次啟動。')
body('③ 在新電腦執行「還原備份」，匯入備份檔即可。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 11. 常見問題
# ══════════════════════════════════════════════════════════════════════════════
h1('11. 常見問題')

h2('Q：App 開啟後顯示「已損壞」或「無法驗證開發者」')
body('執行以下 Terminal 指令，移除 macOS 的隔離旗標，再重新開啟：')
code('xattr -cr /Applications/聖經神學筆記.app')
sp()

h2('Q：App 點了沒有反應，視窗沒有出現')
body('可能有前一次的 App 程序仍在背景執行。請開啟「活動監視器」（Activity Monitor），搜尋「聖經神學筆記」，強制結束該程序後再重新開啟 App。')
sp()

h2('Q：AI 功能顯示錯誤或沒有回應')
bullet('確認已在設定中儲存有效的 API Key（點擊 AI 卡片右上角「設定 API Key」確認）')
bullet('確認目前使用的供應商有足夠的 API 餘額')
bullet('確認網路可正常連線（需連至對應供應商的 API 伺服器）')
bullet('可嘗試切換至其他供應商看是否正常')
sp()

h2('Q：原文（希臘文 / 希伯來文）顯示為方框或亂碼')
body('這是因為系統缺少對應字型。建議安裝以下免費字型：')
bullet('希臘文：Gentium Plus（fonts.sil.org）')
bullet('希伯來文：SBL Hebrew（sbl.org/educational-resources/biblical-fonts）')
body('安裝後重新開啟 App 即可正常顯示。')
sp()

h2('Q：經文查閱面板沒有顯示內容')
bullet('確認筆記已填寫書卷與章節資訊')
bullet('經文資料已內建於 App，無需網路連線，若仍無法顯示請重新啟動 App')
sp()

h2('Q：我的資料存在哪裡？')
body('資料庫位於：')
code('~/Library/Application Support/聖經神學筆記/data/notebook.db')
body('AI Key 等設定位於：')
code('~/Library/Application Support/聖經神學筆記/data/settings.json')
sp()

h2('Q：更新 App 後資料還在嗎？')
body('只要不重新安裝作業系統或手動刪除 Application Support 資料夾，更新 App 後資料完整保留。更新前仍建議先匯出備份以策安全。')
sp()

h2('Q：如何切換使用不同的 AI 供應商？')
body('點擊任一筆記閱讀頁面底部「AI 輔助分析」卡片右上角的「設定 API Key」，選擇目標供應商分頁，點擊「切換至此供應商」即可。也可以在此頁面為多個供應商各自儲存 Key。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
output_path = os.path.expanduser('~/Desktop/聖經神學筆記_本機版使用手冊_v2.docx')
doc.save(output_path)
print(f'已儲存至：{output_path}')
