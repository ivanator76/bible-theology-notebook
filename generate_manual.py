#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate 聖經神學筆記 User Manual as .docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x4A, 0x7E, 0xB5)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(11)
    return p

def note(text):
    """Grey hint/tip paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run('💡 ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def spacer():
    doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════════════
# Cover
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('聖經神學筆記', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

sub = doc.add_paragraph('使用者說明手冊')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

version = doc.add_paragraph('版本 1.0   ·   2025')
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
version.runs[0].font.size = Pt(11)
version.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# Table of Contents (manual)
# ══════════════════════════════════════════════════════════════════════════════
heading1('目錄')
toc_items = [
    ('1', '簡介'),
    ('2', '安裝與啟動'),
    ('3', '主要功能總覽'),
    ('4', '筆記管理'),
    ('5', '標籤系統'),
    ('6', '主題鏈結（Theme Chains）'),
    ('7', '經文快速查閱'),
    ('8', 'AI 輔助分析'),
    ('9', '資源管理'),
    ('10', '備份與還原'),
    ('11', '常見問題'),
]
for num, title_text in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'{num}.  {title_text}')
    run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. 簡介
# ══════════════════════════════════════════════════════════════════════════════
heading1('1. 簡介')
body('聖經神學筆記是一套專為聖經研究者設計的個人筆記系統，整合了以下核心功能：')
bullet('以書卷、章節、節次為索引的聖經研究筆記')
bullet('聖經神學（Biblical Theology）標籤與系統神學（Systematic Theology）標籤')
bullet('主題鏈結：將跨書卷的相關段落串連成研究鏈')
bullet('內建經文查閱：中文（CUV）與英文（WEB）雙語對照')
bullet('AI 輔助分析：相關經文建議、教義連結、延伸研究方向')
bullet('外部資源管理與連結（書籍、文章、影片等）')
bullet('完整備份與還原功能')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 2. 安裝與啟動
# ══════════════════════════════════════════════════════════════════════════════
heading1('2. 安裝與啟動')

heading2('2.1 macOS Electron 版本（本機）')
body('本版本為獨立的 macOS 應用程式，無需安裝伺服器或 Docker。')
bullet('雙擊 .dmg 安裝檔，將「聖經神學筆記.app」拖入「應用程式」資料夾。')
bullet('首次開啟時，macOS 可能顯示「已損壞」或「無法驗證」的警告。')
body('請在 Terminal 執行以下指令移除隔離旗標，再重新開啟 App：')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run('xattr -cr /Applications/聖經神學筆記.app')
run.font.name = 'Courier New'
run.font.size = Pt(10)
run.font.bold = True
bullet('所有資料儲存於本機，路徑：~/Library/Application Support/聖經神學筆記/data/')
spacer()

heading2('2.2 Docker 版本（NAS / 伺服器）')
body('適合部署於 Synology NAS 或任何支援 Docker 的伺服器。')
bullet('在 docker-compose.yml 所在目錄執行：docker-compose up -d')
bullet('預設埠號：3000（可透過環境變數 PORT 修改）')
bullet('資料目錄：./data/（掛載於容器內 /app/backend/data）')
bullet('如需使用 AI 功能，請在 docker-compose.yml 設定 ANTHROPIC_API_KEY 環境變數')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 3. 主要功能總覽
# ══════════════════════════════════════════════════════════════════════════════
heading1('3. 主要功能總覽')
body('應用程式左側為導覽側欄，右側為主要內容區域。側欄包含以下項目：')
bullet('全部筆記：顯示所有筆記，可依標籤篩選')
bullet('依書卷瀏覽：以聖經書卷為分類')
bullet('主題鏈結：管理跨段落的研究主題鏈')
bullet('資源庫：管理外部學習資源')
bullet('標籤管理：管理聖經神學及系統神學標籤')
bullet('備份 / 還原：資料匯出與匯入')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 4. 筆記管理
# ══════════════════════════════════════════════════════════════════════════════
heading1('4. 筆記管理')

heading2('4.1 建立新筆記')
body('點擊右上角「＋ 新增筆記」按鈕，進入筆記編輯頁面。')
bullet('書卷：從下拉選單選擇舊約或新約書卷')
bullet('章節範圍：輸入起始章與結束章（單章留空結束章即可）')
bullet('節次範圍（選填）：輸入起始節與結束節，精確標記段落')
bullet('標題（選填）：筆記標題')
bullet('內容：支援 Markdown 格式（粗體、斜體、標題、列表、引用、程式碼等）')
bullet('聖經神學標籤：標記段落所涉及的聖經神學主題')
bullet('系統神學標籤：標記相關的系統神學教義分類')
note('內容欄支援 Markdown 預覽，輸入時右側會即時顯示排版效果。')
spacer()

heading2('4.2 閱讀筆記')
body('在筆記列表點擊任一筆記即可進入閱讀模式，顯示完整筆記內容、標籤、附加資源及 AI 分析結果。')
spacer()

heading2('4.3 編輯與刪除')
bullet('在閱讀頁面點擊「編輯」按鈕進入編輯模式')
bullet('點擊「刪除」按鈕並確認，即可永久刪除筆記')
note('刪除後無法復原，建議定期備份。')
spacer()

heading2('4.4 筆記搜尋與篩選')
bullet('頂部搜尋欄可依書卷名稱、標題或內容關鍵字搜尋')
bullet('左側標籤面板可點擊標籤進行篩選')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 5. 標籤系統
# ══════════════════════════════════════════════════════════════════════════════
heading1('5. 標籤系統')
body('本系統採用雙層標籤架構，分為聖經神學標籤與系統神學標籤。')

heading2('5.1 聖經神學標籤（Biblical Theology Tags）')
body('用於標記段落在聖經正典敘事中的主題位置，例如：盟約、應許、預表、成全、末世等。')
bullet('前往「標籤管理」頁面新增、編輯或刪除聖經神學標籤')
bullet('標籤可在筆記編輯頁面直接選取或即時新增')
spacer()

heading2('5.2 系統神學標籤（Systematic Theology Tags）')
body('用於標記段落所涉及的系統神學教義，例如：神論、基督論、救贖論、教會論、末世論等。')
bullet('操作方式與聖經神學標籤相同')
bullet('AI 教義連結功能會優先從此標籤庫中提供建議')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 6. 主題鏈結
# ══════════════════════════════════════════════════════════════════════════════
heading1('6. 主題鏈結（Theme Chains）')
body('主題鏈結功能讓您將跨書卷、跨章節的相關聖經段落串連成一條研究鏈，追蹤某個主題在正典中的發展軌跡。')
bullet('點擊「主題鏈結」進入管理頁面')
bullet('點擊「＋ 新增主題鏈」，輸入鏈結名稱與說明')
bullet('在鏈結詳細頁面，可搜尋並加入現有筆記作為節點')
bullet('節點可拖動排序，以反映主題的發展順序')
bullet('從筆記閱讀頁面也可直接將該筆記加入某條主題鏈')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 7. 經文快速查閱
# ══════════════════════════════════════════════════════════════════════════════
heading1('7. 經文快速查閱')
body('在筆記閱讀及編輯頁面，系統會自動根據所記錄的章節範圍，提供內建經文查閱面板。')

heading2('7.1 語言切換')
bullet('中文：顯示中文和合本（CUV）')
bullet('中英：同時顯示中文和英文（WEB）對照')
bullet('英文：顯示英文 Web English Bible（WEB）')
spacer()

heading2('7.2 顯示範圍')
bullet('預設只顯示筆記中指定的節次範圍')
bullet('點擊底部「顯示整章」可展開查看完整章節；再次點擊「只顯示選擇節數」可收回')
spacer()

heading2('7.3 快取機制')
body('已查閱過的章節會暫存於本機 Session，同一章節重複開啟時不需重新下載。')
note('經文資料來源為 bible-api.com，需要網路連線才能首次載入。')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 8. AI 輔助分析
# ══════════════════════════════════════════════════════════════════════════════
heading1('8. AI 輔助分析')
body('AI 功能使用 Anthropic Claude 模型，為您的筆記提供三種輔助分析。')

heading2('8.1 設定 API Key')
body('首次使用前需要設定 Anthropic API Key：')
bullet('點擊筆記閱讀頁面右下角「AI 輔助分析」卡片中的「API Key」按鈕')
bullet('前往 console.anthropic.com 申請 API Key（格式：sk-ant-...）')
bullet('將 Key 貼入輸入欄並點擊「儲存」')
bullet('Key 僅儲存於本機，不會上傳至任何外部伺服器')
note('若已透過環境變數（ANTHROPIC_API_KEY）設定，介面將顯示「由環境變數提供」，無法在介面修改。')
spacer()

heading2('8.2 三種分析功能')

heading3('建議相關經文')
body('AI 根據筆記內容，建議 5–8 段在主題、神學概念或正典脈絡上相關的經文，並附上關聯原因說明。')

heading3('建議教義連結')
body('AI 分析此段落如何貢獻於系統神學各教義，優先從您現有的系統神學標籤中選取，並說明具體的神學連結點。')

heading3('延伸研究方向')
body('AI 建議 3–5 個值得深入探索的研究方向，包含正典關聯、歷史文化背景、跨文本比較及重要神學問題。')
spacer()

heading2('8.3 使用方式')
bullet('在筆記閱讀頁面滾動至底部，找到「AI 輔助分析」卡片')
bullet('點擊三個功能按鈕之一，等待 AI 回應（約 10–30 秒）')
bullet('結果以 Markdown 格式顯示，可直接閱讀')
bullet('同一筆記的分析結果暫存於本次工作階段，不會永久儲存')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 9. 資源管理
# ══════════════════════════════════════════════════════════════════════════════
heading1('9. 資源管理')
body('資源庫讓您管理與聖經研究相關的外部資源，包括書籍、文章、影片、講道、網頁等。')
bullet('點擊「資源庫」進入管理頁面')
bullet('點擊「＋ 新增資源」，填寫資源名稱、類型、作者、連結等資訊')
bullet('在筆記閱讀頁面，可將資源連結至特定筆記')
bullet('資源可加上頁碼或段落備註，方便日後查閱')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 10. 備份與還原
# ══════════════════════════════════════════════════════════════════════════════
heading1('10. 備份與還原')

heading2('10.1 匯出備份')
bullet('點擊側欄「備份」')
bullet('點擊「匯出備份」，下載 JSON 格式的完整備份檔')
bullet('備份包含所有筆記、標籤、主題鏈結及資源資料')
note('建議定期（每週或每月）手動備份，並將備份檔儲存至雲端或外部硬碟。')
spacer()

heading2('10.2 還原備份')
bullet('點擊「還原備份」，選取先前匯出的 JSON 備份檔')
bullet('確認還原後，系統將以備份資料取代現有資料')
note('還原操作將覆蓋目前所有資料，請確認後再執行。')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 11. 常見問題
# ══════════════════════════════════════════════════════════════════════════════
heading1('11. 常見問題')

heading2('Q：App 開啟時顯示「已損壞」或「無法驗證開發者」')
body('這是 macOS Gatekeeper 的安全機制，並非 App 真的損壞。請在 Terminal 執行：')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run('xattr -cr /Applications/聖經神學筆記.app')
run.font.name = 'Courier New'
run.font.size = Pt(10)
run.font.bold = True
body('完成後重新開啟 App 即可。')
spacer()

heading2('Q：AI 功能無回應或顯示錯誤')
bullet('確認已設定有效的 Anthropic API Key')
bullet('確認網路可連線至 api.anthropic.com')
bullet('API Key 可於 console.anthropic.com 確認是否有效及餘額')
spacer()

heading2('Q：經文無法載入')
bullet('確認網路連線正常（經文資料來源：bible-api.com）')
bullet('若曾成功載入，同一章節的資料會暫存於本次 Session')
spacer()

heading2('Q：資料儲存在哪裡？')
body('macOS Electron 版本：')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run('~/Library/Application Support/聖經神學筆記/data/notebook.db')
run.font.name = 'Courier New'
run.font.size = Pt(10)
body('Docker 版本：掛載目錄下的 data/notebook.db')
spacer()

heading2('Q：如何將資料從一台 Mac 移到另一台？')
bullet('在舊電腦使用「備份」功能匯出 JSON 備份檔')
bullet('在新電腦安裝 App 後，使用「還原」功能匯入備份檔')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
output_path = os.path.expanduser('~/Desktop/聖經神學筆記_使用者手冊.docx')
doc.save(output_path)
print(f'已儲存至：{output_path}')
