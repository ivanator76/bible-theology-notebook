#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate 聖經神學筆記 NAS 版使用手冊 .docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ── Page margins (A4) ─────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Inches(8.27)
section.page_height   = Inches(11.69)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
ACCENT  = RGBColor(0x2C, 0x5F, 0x8A)
ACCENT2 = RGBColor(0x4A, 0x7E, 0xB5)
MUTED   = RGBColor(0x55, 0x55, 0x55)
GREY    = RGBColor(0x88, 0x88, 0x88)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = ACCENT
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = ACCENT2
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run('提示：' + text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = MUTED
    return p

def sp():
    doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════════════════════
t = doc.add_heading('聖經神學筆記', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

p = doc.add_paragraph('NAS 版　使用手冊')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(14)
p.runs[0].font.color.rgb = MUTED

p2 = doc.add_paragraph('版本 2.0')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.runs[0].font.size = Pt(11)
p2.runs[0].font.color.rgb = GREY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 目錄
# ══════════════════════════════════════════════════════════════════════════════
h1('目錄')
for item in [
    '1.  關於 NAS 版',
    '2.  開始使用',
    '3.  介面說明',
    '4.  建立與管理筆記',
    '5.  讀經模式',
    '6.  標籤系統',
    '7.  主題鏈結',
    '8.  經文查閱',
    '9.  匯出研究文件',
    '10. AI 輔助分析',
    '11. 資源庫',
    '12. 備份與還原',
    '13. 從外網存取（Tailscale）',
    '14. 常見問題',
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(item)
    run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. 關於 NAS 版
# ══════════════════════════════════════════════════════════════════════════════
h1('1. 關於 NAS 版')
body('NAS 版聖經神學筆記部署於您家中的 NAS（網路儲存裝置）上，透過瀏覽器使用，無需在電腦或行動裝置上安裝任何應用程式。')
sp()

body('主要特點：')
bullet('任何裝置皆可使用：電腦（Windows / macOS / Linux）、iPad、iPhone、Android，只需瀏覽器')
bullet('資料存放於 NAS 本機，完全私人，不依賴任何雲端服務')
bullet('內建四種語言聖經（中文和合本、英文 WEB、希臘文、希伯來文），完全離線可用')
bullet('支援三大 AI 供應商（Anthropic / OpenAI / Google）')
bullet('家中區域網路直接存取，也可透過 Tailscale 從外網連線')
sp()

body('存取網址（家中 Wi-Fi）：')
code('http://192.168.50.8:3000')
tip('實際 IP 視您的網路設定而定，請向管理員確認。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 2. 開始使用
# ══════════════════════════════════════════════════════════════════════════════
h1('2. 開始使用')

h2('2.1 連線方式')
body('在與 NAS 同一個 Wi-Fi 網路下，開啟任意瀏覽器，輸入網址即可：')
code('http://192.168.50.8:3000')
body('建議瀏覽器：Chrome、Safari、Firefox、Edge（任一現代瀏覽器皆可）。')
sp()
tip('可以將此網址加入瀏覽器書籤，或在 iPhone / iPad 上「加入主畫面」，使用起來就像 App 一樣方便。')
sp()

h2('2.2 將網頁加入主畫面（iPhone / iPad）')
body('① 在 Safari 開啟 http://192.168.50.8:3000')
body('② 點擊底部分享按鈕（□↑）')
body('③ 選擇「加入主畫面」')
body('④ 名稱填「聖經筆記」，點擊「新增」')
body('完成後主畫面會出現圖示，點擊後直接以全螢幕模式開啟。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 3. 介面說明
# ══════════════════════════════════════════════════════════════════════════════
h1('3. 介面說明')
body('主介面分為左右兩個區域：')
sp()

h2('左側：導覽側欄')
bullet('全部筆記：列出所有筆記，可依標籤篩選')
bullet('讀經模式：逐章閱讀經文，已有筆記的經節會直接標示出來')
bullet('依書卷瀏覽：以聖經書卷分類瀏覽')
bullet('主題鏈結：跨段落的主題研究鏈管理')
bullet('資源庫：外部學習資源管理')
bullet('標籤管理：建立與管理聖經神學及系統神學標籤')
bullet('備份 / 還原：資料匯出與匯入')
sp()

h2('右側：主內容區域')
body('依據左側所選項目，顯示筆記列表、筆記詳情、編輯表單或管理頁面。')
sp()

h2('頂部搜尋列')
body('按 Ctrl+K（Windows）或 ⌘K（Mac）開啟全域搜尋，輸入兩個字以上即開始搜尋。'
     '搜尋範圍涵蓋四類內容，結果會分類顯示：')
bullet('筆記：標題、內文、經文範圍（含書卷中英文名）')
bullet('教義註解：教義連結裡填寫的說明文字')
bullet('追蹤鏈：名稱與說明')
bullet('資料：外部資料的標題、作者、出處、摘要、網址')
body('搜尋結果會顯示命中處的上下文片段，關鍵字以高亮標示，'
     '讓您不必逐一點開就能判斷哪一筆才是要找的。'
     '按上下鍵移動選取、Enter 開啟、Esc 關閉。')
sp()
tip('搜尋在伺服器端以全文索引進行，筆記量增加也不會變慢；'
    '新增或修改內容後索引會自動更新。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 4. 建立與管理筆記
# ══════════════════════════════════════════════════════════════════════════════
h1('4. 建立與管理筆記')

h2('4.1 新增筆記')
body('點擊右上角「＋ 新增筆記」，進入編輯頁面，填寫以下欄位：')
sp()

h3('書卷與章節')
bullet('書卷：從下拉選單選取（舊約 / 新約）')
bullet('起始章 / 結束章：輸入章次（單章筆記只填起始章即可）')
bullet('起始節 / 結束節（選填）：精確標記段落節次，例如 1:1–18')
sp()

h3('標題與內容')
bullet('標題（選填）：筆記的簡短名稱')
bullet('內容：支援 Markdown 語法，輸入時右側即時預覽排版')
sp()

h3('常用 Markdown 語法')
for l in [
    '# 大標題    ## 中標題    ### 小標題',
    '**粗體**    *斜體*    > 引用區塊',
    '- 項目列表    1. 編號列表',
    '`行內程式碼`',
]:
    rr = doc.add_paragraph()
    rr.paragraph_format.left_indent = Cm(1)
    run = rr.add_run(l)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
sp()

h3('標籤')
bullet('聖經神學標籤：點擊下拉選取，或直接輸入新標籤名稱後按 Enter 建立')
bullet('系統神學標籤：操作方式相同')
sp()

h2('4.2 閱讀筆記')
body('在筆記列表點擊任一筆記進入閱讀模式，顯示：')
bullet('完整筆記內容（Markdown 渲染後）')
bullet('所有標籤、所屬主題鏈結、連結的外部資源')
bullet('下方的經文查閱面板')
bullet('底部的 AI 輔助分析面板')
sp()

h2('4.3 編輯與刪除')
bullet('閱讀頁面右上角點擊「編輯」進入編輯模式')
bullet('編輯中按 Ctrl+S（Windows）或 ⌘S（Mac）可直接儲存')
bullet('點擊「刪除」並在確認對話框確認，即永久刪除')
tip('刪除後無法復原，建議定期使用「備份」功能匯出資料。')
sp()

h2('4.4 草稿自動保護')
body('編輯筆記時，內容每隔約兩秒會自動存成草稿（保存於瀏覽器本機），'
     '編輯畫面上方會顯示「草稿已自動儲存」與時間。')
sp()
body('若分頁意外關閉、瀏覽器當機或連線中斷，下次開啟同一則筆記時，'
     '上方會出現「找到自動儲存草稿」的提示：')
bullet('點「還原草稿」：把尚未儲存的內容救回來')
bullet('點「捨棄」：丟棄草稿，使用已儲存的版本')
sp()
tip('正常按下「儲存」後草稿會自動清除；尚未儲存的新筆記草稿同樣會被保留。')
sp()

h2('4.5 修訂歷史')
body('每次儲存有變動的筆記時，系統會自動保留一份修改前的快照。'
     '在筆記閱讀頁面的「修訂歷史」卡片中：')
bullet('點「查看歷史」列出所有舊版本（時間、標題、字數）')
bullet('點「預覽」查看該版本的完整內容')
bullet('點「還原」把筆記回復到該版本')
sp()
tip('還原前系統會先把目前的內容也存成一個版本，因此還原後仍可再回復回來。'
    '內容完全沒有變動的儲存不會產生重複版本。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 5. 讀經模式
# ══════════════════════════════════════════════════════════════════════════════
h1('5. 讀經模式')
body('一般的筆記系統是「先想到筆記，再找經文」。讀經模式是反過來的：'
     '打開一卷書逐章閱讀，過去寫過的筆記會自動出現在對應的經節旁邊。')
sp()

h2('5.1 開啟方式')
body('左側導覽側欄點擊「讀經模式 Reading」。')
sp()

h2('5.2 翻閱經文')
bullet('書卷下拉選單：選擇要閱讀的書卷')
bullet('上一章 / 下一章：在同一卷書內翻章')
bullet('章節下拉選單：直接跳到指定章')
body('經文以中英對照顯示（和合本 + WEB 英文譯本）。')
sp()

h2('5.3 已有筆記的經節')
bullet('該節會以較明顯的樣式標示')
bullet('節文下方出現筆記標題的小標籤，點擊即可跳到該則筆記')
bullet('一節若對應多則筆記，會全部列出')
body('頁面頂端顯示「本章 N 節已有筆記 · 共 M 節」，'
     '可快速看出這一章的研讀密度。')
sp()
tip('涵蓋範圍會自動計算：整章筆記、單節筆記、跨章筆記都會正確對應到所屬的每一節。')
sp()

h2('5.4 邊讀邊寫')
bullet('頁面右上角「本章新增筆記」：建立一則涵蓋整章的筆記')
bullet('每一節右側的「＋」：針對該節新增筆記')
body('由讀經模式進入新增筆記時，書卷、章、節會自動帶入，不需重新選擇。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 6. 標籤系統
# ══════════════════════════════════════════════════════════════════════════════
h1('6. 標籤系統')
body('本系統使用兩種獨立的標籤，讓您從不同角度標記同一段聖經文本。')
sp()

h2('6.1 聖經神學標籤（BT Tags）')
body('用於標記段落在聖經正典敘事中的主題位置與神學功能，例如：')
bullet('盟約、應許、預表、成全、王權、聖殿、出埃及、新創造……')
body('管理路徑：側欄「標籤管理」→「聖經神學標籤」')
sp()

h2('6.2 系統神學標籤（ST Tags）')
body('用於標記段落所涉及的系統神學教義分類，例如：')
bullet('神論、基督論、聖靈論、人論、罪論、救贖論、教會論、末世論……')
body('管理路徑：側欄「標籤管理」→「系統神學標籤」')
sp()

h2('6.3 透過標籤篩選筆記')
body('在「全部筆記」頁面，點擊左側任一標籤，即可篩選出包含該標籤的所有筆記。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 6. 主題鏈結
# ══════════════════════════════════════════════════════════════════════════════
h1('7. 主題鏈結（Theme Chains）')
body('主題鏈結讓您將跨書卷、跨章節的相關段落串成一條研究鏈，追蹤某個神學主題在正典中的發展軌跡。')
sp()

h2('7.1 建立主題鏈')
bullet('點擊側欄「主題鏈結」')
bullet('點擊「＋ 新增」，輸入鏈結名稱與說明')
sp()

h2('7.2 加入節點')
bullet('進入鏈結詳細頁面，點擊「加入筆記」')
bullet('搜尋並選取已有的筆記作為節點')
bullet('可拖動節點排序，以反映主題的發展順序')
sp()

h2('7.3 從筆記加入鏈結')
body('在筆記閱讀頁面，可直接將該筆記加入某條已有的主題鏈，無需切換頁面。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 7. 經文查閱
# ══════════════════════════════════════════════════════════════════════════════
h1('8. 經文查閱')
body('在筆記閱讀及編輯頁面，下方會顯示對應章節的「經文參考」面板。所有經文資料已離線內建，無需網路連線。')
sp()

h2('8.1 語言切換')
body('面板頂部有四個語言按鈕：')

table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '按鈕'
hdr[1].text = '說明'
rows_data = [
    ('中文', '顯示中文和合本（CUV）'),
    ('中英', '同時顯示中文與英文（WEB）雙語對照'),
    ('English', '僅顯示英文 World English Bible（WEB）'),
    ('原文', '新約顯示希臘文（Tischendorf 8th ed.）；舊約顯示希伯來文（Westminster Leningrad Codex）'),
]
for i, (btn, desc) in enumerate(rows_data):
    row = table.rows[i+1]
    row.cells[0].text = btn
    row.cells[1].text = desc
sp()

h2('8.2 原文顯示')
bullet('新約（馬太福音至啟示錄）：顯示帶有氣息符號與重音的古希臘文（Tischendorf 第八版）')
bullet('舊約（創世記至瑪拉基書）：顯示帶有母音符號（尼庫達）的希伯來文（Westminster Leningrad Codex）')
bullet('希伯來文採由右至左（RTL）排列，為正確閱讀方向')
sp()

h2('8.3 顯示範圍')
bullet('預設顯示筆記所指定的節次範圍，指定節次會以金色色塊標示')
bullet('點擊面板底部「顯示整章」，可展開查看完整章節')
bullet('再次點擊「只顯示選擇節數」，可收回至指定範圍')
sp()

h2('8.4 調整字體大小')
body('面板右上角有字體大小控制器（A− / 數字 / A+），可自由調整顯示大小（範圍：11–22px）。調整後的設定會記憶於瀏覽器，下次開啟時保持不變。')
sp()

h2('8.5 摺疊面板')
body('點擊面板標題列的箭頭圖示可摺疊 / 展開面板，不影響筆記內容。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 9. 匯出研究文件
# ══════════════════════════════════════════════════════════════════════════════
h1('9. 匯出研究文件')
body('把累積的研究整批編譯成可交付的文件——備講道、寫查經材料、'
     '或與人討論神學時使用。')
sp()

h2('9.1 可匯出的兩種範圍')
bullet('一條主題鏈：在主題鏈結頁面選定某條鏈後，點「匯出研究」')
bullet('一個教義：在教義頁面選定某個教義後，點「匯出研究」')
sp()

h2('9.2 三種格式')

table_fmt = doc.add_table(rows=4, cols=2)
table_fmt.style = 'Table Grid'
hdr_fmt = table_fmt.rows[0].cells
hdr_fmt[0].text = '格式'
hdr_fmt[1].text = '適用場景'
for i, (fmt, use) in enumerate([
    ('Markdown', '貼進其他筆記軟體、版本控管、再加工'),
    ('Word', '需要進一步排版、加註、與人協作'),
    ('PDF', '直接列印或分發，版面固定'),
]):
    row = table_fmt.rows[i + 1]
    row.cells[0].text = fmt
    row.cells[1].text = use
sp()

h2('9.3 文件內容')
body('匯出的文件會依正典順序排列所有相關筆記，每則包含：')
bullet('經文出處與筆記標題')
bullet('經文全文（和合本 + WEB 英文對照）')
bullet('筆記內容')
bullet('該筆記的教義註解')
bullet('該筆記連結的參考資料清單（作者、出處、頁碼、網址）')
sp()
tip('PDF 需要伺服器具備中文字型。Docker 映像已內建 Noto CJK 字型；'
    '若自行部署於其他環境而 PDF 匯出報錯，可用 PDF_FONT_PATH 環境變數指定字型檔路徑。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 10. AI 輔助分析
# ══════════════════════════════════════════════════════════════════════════════
h1('10. AI 輔助分析')
body('AI 功能支援四家 AI 服務供應商，提供三種智慧分析功能。'
     '分析結果會保存於資料庫，並可一鍵寫入知識庫。使用前需設定 API Key 及網路連線。')
sp()

h2('10.1 支援的 AI 供應商')

table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = '供應商'
hdr2[1].text = '使用的模型'
hdr2[2].text = 'API Key 格式'
providers = [
    ('Anthropic', 'Claude Sonnet', 'sk-ant-...'),
    ('OpenAI', 'GPT-4o', 'sk-...'),
    ('Google', 'Gemini 2.0 Flash', 'AIza...'),
    ('OpenRouter', '數百種模型可自由選擇', 'sk-or-v1-...'),
]
for i, (name, model, fmt) in enumerate(providers):
    row = table2.rows[i+1]
    row.cells[0].text = name
    row.cells[1].text = model
    row.cells[2].text = fmt
sp()
body('OpenRouter 是一個聚合平台：用一把 Key 就能呼叫上百種不同廠牌的模型'
     '（Anthropic、OpenAI、Google、DeepSeek、Llama 等），並可隨時切換比較。'
     '適合想試不同模型、或想集中管理帳單的情況。')
sp()

h2('10.2 設定 API Key')
body('① 在任一筆記閱讀頁面，找到底部「AI 輔助分析」卡片。')
body('② 點擊卡片右上角「API Key」按鈕，開啟「AI 服務設定」視窗。')
body('③ 選擇您要使用的供應商分頁（Anthropic / OpenAI / Google / OpenRouter）。')
body('④ 取得對應的 API Key，貼入輸入欄，點擊「儲存」，即自動切換為使用該供應商。')
body('⑤ 若已設定多個供應商，點擊「切換為使用此服務」可在不同 AI 服務間切換。')
sp()
body('已存有 Key 的分頁名稱旁會顯示小圓點，目前使用中的分頁會標示「✓ 目前使用中」。'
     '要移除某個 Key，切到該分頁點「移除 Key」；'
     '若移除的是使用中的供應商，系統會自動改用其他還有 Key 的供應商。')
sp()

h3('取得 API Key 的網址')
bullet('Anthropic：console.anthropic.com')
bullet('OpenAI：platform.openai.com/api-keys')
bullet('Google：aistudio.google.com')
bullet('OpenRouter：openrouter.ai/keys')
sp()
tip('API Key 儲存於 NAS 上的 data/settings.json，不會上傳至任何外部伺服器。')
sp()

h3('OpenRouter：選擇要使用的模型')
body('選擇 OpenRouter 分頁時，會多出一個「使用的模型」欄位：')
bullet('點擊欄位會出現完整模型清單（從 OpenRouter 即時取得，數百種可選）')
bullet('也可直接輸入模型代號，例如 anthropic/claude-sonnet-4.5、google/gemini-2.5-flash')
bullet('留空則使用預設模型 anthropic/claude-sonnet-4.5')
bullet('修改後離開欄位即自動儲存')
body('設定完成後，AI 輔助分析面板的標題旁會顯示目前的供應商與模型，'
     '例如「OpenRouter · gemini-2.5-flash」。')
sp()
tip('若模型清單載入失敗（例如暫時連不上 OpenRouter），'
    '欄位旁會出現「模型清單載入失敗，重試」，此時仍可手動輸入模型代號。')
sp()

h2('10.3 三種分析功能')

h3('① 建議相關經文')
body('AI 根據筆記內容、標題及標籤，建議 5–8 段在主題、神學概念或正典脈絡上與此段落相關的經文，並附上關聯原因說明。')
sp()

h3('② 建議教義連結')
body('AI 分析此段落如何貢獻於系統神學各教義，優先從您現有的系統神學標籤中選取最相關的 2–4 個教義，說明具體的神學連結點。')
sp()

h3('③ 延伸研究方向')
body('AI 建議 3–5 個值得深入探索的研究方向，涵蓋正典關聯、歷史文化背景、跨文本比較及重要神學問題。')
sp()

h2('10.4 使用流程')
body('① 開啟任一筆記的閱讀頁面，滾動至底部的「AI 輔助分析」卡片。')
body('② 卡片標題旁會顯示目前使用的 AI 供應商與模型。')
body('③ 點擊三個功能按鈕之一，等待 AI 回應（通常 10–30 秒）。')
body('④ 結果會以逐條建議的形式顯示，每條建議都可以「採納」或「忽略」。')
sp()

h2('10.5 把建議直接寫入知識庫')
body('AI 的每一條建議下方都有「採納」與「忽略」兩個按鈕：')

table_adopt = doc.add_table(rows=4, cols=2)
table_adopt.style = 'Table Grid'
hdr_a = table_adopt.rows[0].cells
hdr_a[0].text = '建議類型'
hdr_a[1].text = '點「採納」的效果'
for i, (kind, effect) in enumerate([
    ('建議教義連結', '直接建立教義連結，把 AI 寫的註解存入該筆記'),
    ('建議相關經文', '若該段經文已有筆記，直接建立交叉引用'),
    ('延伸研究方向', '標記為已採納（僅作為閱讀紀錄）'),
]):
    row = table_adopt.rows[i + 1]
    row.cells[0].text = kind
    row.cells[1].text = effect
sp()
bullet('建議相關經文時，若該段落已有您的筆記，旁邊會出現「已有筆記」按鈕可直接跳過去查看')
bullet('若該段落還沒有筆記，「採納」會停用並提示先到讀經模式建立筆記')
bullet('已處理的建議會標示「已採納」或「已忽略」')
sp()

h2('10.6 分析結果會被保存')
body('分析結果存於資料庫，離開頁面再回來仍在，'
     '不需重跑、也不會重複消耗 API 費用。')
bullet('已有結果的功能按鈕旁會出現小圓點標記')
bullet('點該按鈕直接顯示上次的結果')
bullet('需要重跑時，點結果區的「重新分析」')
bullet('即使之後移除了 API Key，已保存的分析仍可查看')
sp()
tip('AI 分析結果會一併包含在匯出備份中。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 9. 資源庫
# ══════════════════════════════════════════════════════════════════════════════
h1('11. 資源庫')
body('資源庫讓您集中管理與聖經研究相關的外部資源，並可連結至特定筆記。')
sp()

h2('11.1 新增資源')
bullet('點擊側欄「資源庫」，再點擊「＋ 新增資源」')
bullet('填寫資源名稱、類型（書籍 / 文章 / 影片 / 講道 / 網頁等）')
bullet('可輸入作者、出版年份、網址連結等資訊')
sp()

h2('11.2 連結資源至筆記')
bullet('在筆記閱讀或編輯頁面，找到「相關資源」區塊')
bullet('點擊「連結資源」，從資源庫中選取並加入')
bullet('可加上頁碼或段落備註，方便日後查閱')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 10. 備份與還原
# ══════════════════════════════════════════════════════════════════════════════
h1('12. 備份與還原')
body('強烈建議定期備份，以防止資料遺失。')
sp()

h2('12.1 匯出備份')
bullet('點擊側欄「備份 / 還原」')
bullet('點擊「匯出備份」，系統將下載一份 JSON 格式的完整備份檔')
bullet('備份檔包含所有筆記、標籤、主題鏈結、教義連結、交叉引用、資源資料、'
       '筆記修訂歷史與 AI 分析結果')
tip('建議每週或每月定期備份，並將備份檔存至其他安全位置。')
sp()

h2('12.2 還原備份')
bullet('點擊「還原備份」，選取先前匯出的 .json 備份檔')
bullet('系統將以備份資料完全取代目前所有資料')
tip('還原操作不可逆，執行前請確認已了解目前資料將被覆蓋。')
sp()

h2('12.3 資料庫直接備份')
body('更完整的備份方式是直接複製 NAS 上的 SQLite 資料庫檔案：')
code('data/notebook.db')
body('此檔案位於 NAS 的掛載目錄下，可定期複製備份。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 11. 從外網存取（Tailscale）
# ══════════════════════════════════════════════════════════════════════════════
h1('13. 從外網存取（Tailscale）')
body('若您不在家中 Wi-Fi 網路，可透過 Tailscale 從任何地方安全連線至 NAS。')
sp()

h2('13.1 什麼是 Tailscale？')
body('Tailscale 是一套免費的個人 VPN 服務，在您的裝置和 NAS 之間建立私人加密通道，無需開放路由器 port，安全性極高。')
sp()

h2('13.2 設定步驟')
body('① 在 NAS 的 Synology 套件中心搜尋並安裝「Tailscale」。')
body('② 登入 Tailscale 帳號（可用 Google 帳號登入）。')
body('③ 在您的裝置（電腦 / iPhone / iPad）安裝 Tailscale App：')
bullet('電腦：tailscale.com/download')
bullet('iPhone / iPad：App Store 搜尋「Tailscale」')
body('④ 在所有裝置上登入同一個 Tailscale 帳號。')
body('⑤ 完成後，Tailscale 會為 NAS 分配一個固定的虛擬 IP（如 100.x.x.x）。')
sp()

h2('13.3 外網存取方式')
body('開啟 Tailscale 後，在瀏覽器輸入 NAS 的 Tailscale IP：')
code('http://100.x.x.x:3000')
body('（實際 IP 請在 Tailscale 管理介面或 App 中查看）')
sp()
tip('Tailscale 免費方案支援最多 3 台裝置，個人使用完全足夠。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# 12. 常見問題
# ══════════════════════════════════════════════════════════════════════════════
h1('14. 常見問題')

h2('Q：在家中無法開啟網頁')
bullet('確認裝置與 NAS 在同一個 Wi-Fi 網路')
bullet('確認 NAS 上的 Docker 容器正在執行')
bullet('嘗試直接輸入 IP 位址，避免 DNS 問題')
sp()

h2('Q：外網無法連線')
bullet('確認 NAS 和您的裝置都已安裝並啟動 Tailscale')
bullet('確認兩者使用同一個 Tailscale 帳號')
bullet('在 Tailscale 管理介面確認 NAS 顯示為「Connected」')
sp()

h2('Q：AI 功能顯示錯誤或沒有回應')
bullet('確認已在設定中儲存有效的 API Key')
bullet('確認目前使用的供應商有足夠的 API 餘額')
bullet('確認網路可正常連線（AI 需要連至供應商 API 伺服器）')
bullet('可嘗試切換至其他供應商')
sp()

h2('Q：原文（希臘文 / 希伯來文）顯示為方框或亂碼')
body('這是因為裝置缺少對應字型。建議安裝以下免費字型：')
bullet('希臘文：Gentium Plus（fonts.sil.org）')
bullet('希伯來文：SBL Hebrew（sbl.org）')
body('安裝後重新整理頁面即可正常顯示。')
sp()

h2('Q：資料儲存在哪裡？')
body('所有筆記資料儲存於 NAS 上的掛載目錄：')
code('data/notebook.db')
body('AI Key 等設定儲存於：')
code('data/settings.json')
body('只要不刪除此目錄，重新部署容器後資料完整保留。')
sp()

h2('Q：如何切換使用不同的 AI 供應商？')
body('點擊任一筆記閱讀頁面底部「AI 輔助分析」卡片右上角的「API Key」，'
     '選擇目標供應商分頁，點擊「切換為使用此服務」即可。'
     '若使用 OpenRouter，還可在同一個分頁的「使用的模型」欄位切換模型。')
sp()

h2('Q：不小心關掉分頁，寫到一半的筆記還在嗎？')
body('在。編輯中的內容每約兩秒會自動存成草稿，'
     '重新開啟同一則筆記時會出現「找到自動儲存草稿」提示，點「還原草稿」即可救回。'
     '已儲存過的筆記若想回到更早的版本，可用閱讀頁面的「修訂歷史」還原。')
sp()

h2('Q：可以把整條主題鏈的研究成果印出來嗎？')
body('可以。在主題鏈結或教義頁面選定項目後，點「匯出研究」，'
     '可選擇 Markdown、Word 或 PDF 三種格式。'
     '文件會依正典順序排列，包含經文全文、筆記內容、教義註解與參考資料清單。')
sp()

h2('Q：可以多人同時使用嗎？')
body('可以。NAS 版為網頁服務，同一網路下的多台裝置可同時存取，但所有人共用同一份資料庫（筆記、標籤等），不支援多使用者帳號隔離。')
sp()

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
output_path = os.path.expanduser('~/Desktop/聖經神學筆記_NAS版使用手冊_v2.docx')
doc.save(output_path)
print(f'已儲存至：{output_path}')
